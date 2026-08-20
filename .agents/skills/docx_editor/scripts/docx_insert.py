#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
docx_insert.py - Word 文档段落插入工具
在文档的指定位置（按关键词定位或按段落序号定位）插入新的段落或文本内容，
新段落自动继承上下文现有段落的样式。
用法:
  python docx_insert.py input.docx --after-text "第二章结论" --content "这是新增的总结段落" --output out.docx
  python docx_insert.py input.docx --after-index 5 --content "在第5段后插入的新段落" --output out.docx
  python docx_insert.py input.docx --at-end --content "文末追加的段落" --output out.docx
"""

import sys
import shutil
import argparse
from pathlib import Path

# 确保 Windows 终端 UTF-8 编码正常输出
if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8")
if hasattr(sys.stderr, "reconfigure"):
    sys.stderr.reconfigure(encoding="utf-8")

try:
    import docx
    from docx.shared import Pt
    from docx.oxml.ns import qn
    from copy import deepcopy
except ImportError:
    print("错误: 缺少 python-docx 依赖，请运行: uv pip install --system python-docx 或 pip install python-docx")
    sys.exit(1)


def backup_file(file_path: str) -> str:
    """在修改前自动创建 .bak.docx 备份副本"""
    p = Path(file_path)
    if not p.exists():
        return ""
    bak_path = p.with_suffix(".bak.docx")
    shutil.copy2(str(p), str(bak_path))
    return str(bak_path)


def _insert_paragraph_after(reference_paragraph, text: str, inherit_style: bool = True):
    """
    在指定段落后面插入新段落。
    如果 inherit_style 为 True，新段落会继承参考段落的样式。
    """
    new_p_element = deepcopy(reference_paragraph._element)
    # 清空文本内容，保留样式属性
    for child in list(new_p_element):
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag == 'r':  # 删除所有 Run（文字内容）
            new_p_element.remove(child)
    
    # 插入到参考段落之后
    reference_paragraph._element.addnext(new_p_element)
    
    # 获取新创建的段落对象
    new_paragraph = docx.text.paragraph.Paragraph(new_p_element, reference_paragraph._element.getparent())
    
    # 添加文本内容
    run = new_paragraph.add_run(text)
    
    # 如果原段落有 run 且有字体设置，继承字体
    if inherit_style and reference_paragraph.runs:
        ref_run = reference_paragraph.runs[0]
        if ref_run.font.name:
            run.font.name = ref_run.font.name
        if ref_run.font.size:
            run.font.size = ref_run.font.size
        # 继承中文字体
        try:
            rPr = ref_run._element.find(qn('w:rPr'))
            if rPr is not None:
                rFonts = rPr.find(qn('w:rFonts'))
                if rFonts is not None:
                    ea_font = rFonts.get(qn('w:eastAsia'))
                    if ea_font:
                        new_rPr = run._element.get_or_add_rPr()
                        new_rFonts = new_rPr.find(qn('w:rFonts'))
                        if new_rFonts is None:
                            from lxml import etree
                            new_rFonts = etree.SubElement(new_rPr, qn('w:rFonts'))
                        new_rFonts.set(qn('w:eastAsia'), ea_font)
        except Exception:
            pass  # 继承中文字体失败不影响核心功能
    
    return new_paragraph


def insert_content(file_path: str, content: str, output_path: str = None,
                   after_text: str = None, after_index: int = None, at_end: bool = False):
    """
    在文档指定位置插入新段落。
    定位方式（三选一）：
    - after_text: 在包含该文本的第一个段落后面插入
    - after_index: 在第 N 个段落后面插入（0-indexed）
    - at_end: 在文档末尾追加
    """
    doc = docx.Document(file_path)
    save_path = output_path or file_path
    
    paragraphs = doc.paragraphs
    if not paragraphs:
        return "错误: 文档中没有任何段落"
    
    target_paragraph = None
    location_desc = ""
    
    if after_text:
        # 按关键词定位
        for p in paragraphs:
            if after_text in p.text:
                target_paragraph = p
                location_desc = f"包含「{after_text}」的段落之后"
                break
        if target_paragraph is None:
            return f"错误: 文档中未找到包含「{after_text}」的段落"
            
    elif after_index is not None:
        # 按序号定位（0-indexed）
        if after_index < 0 or after_index >= len(paragraphs):
            return f"错误: 段落序号 {after_index} 超出范围（共 {len(paragraphs)} 个段落，有效范围 0~{len(paragraphs)-1}）"
        target_paragraph = paragraphs[after_index]
        location_desc = f"第 {after_index + 1} 个段落之后"
        
    elif at_end:
        # 文档末尾
        target_paragraph = paragraphs[-1]
        location_desc = "文档末尾"
    else:
        return "错误: 必须指定插入位置，使用 --after-text、--after-index 或 --at-end"
    
    # 自动备份最终将被覆盖的文件
    bak_msg = ""
    if Path(save_path).exists():
        bak = backup_file(save_path)
        if bak:
            bak_msg = f"（原文件已备份至: {bak}）"
    
    # 支持多段落插入（同时兼容字面量 \n、真实换行符 \n 以及 \r\n）
    normalized_content = content.replace("\r\n", "\n").replace("\\n", "\n")
    content_parts = normalized_content.split("\n")
    inserted_count = 0
    current_ref = target_paragraph
    for part in content_parts:
        part = part.strip()
        if not part:
            continue
        new_p = _insert_paragraph_after(current_ref, part, inherit_style=True)
        current_ref = new_p
        inserted_count += 1
    
    Path(save_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(save_path)
    return (f"成功在{location_desc}插入了 {inserted_count} 个新段落，"
            f"已保存至: {save_path}{bak_msg}")


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Word 文档段落插入工具")
    parser.add_argument("file_path", help="输入文档路径")
    
    # 定位方式（三选一）
    location = parser.add_mutually_exclusive_group(required=True)
    location.add_argument("--after-text", help="在包含指定文本的段落后面插入")
    location.add_argument("--after-index", type=int, help="在第 N 个段落后面插入（0-indexed）")
    location.add_argument("--at-end", action="store_true", help="在文档末尾追加")
    
    parser.add_argument("--content", required=True, help="要插入的文本内容（多段用 \\n 分隔）")
    parser.add_argument("--output", "-o", help="输出路径（可选，默认覆盖原文件）")
    args = parser.parse_args()
    
    msg = insert_content(
        args.file_path, args.content, args.output,
        after_text=args.after_text,
        after_index=args.after_index,
        at_end=args.at_end
    )
    print(msg)

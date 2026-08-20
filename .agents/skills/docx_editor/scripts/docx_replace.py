#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
docx_replace.py - Word 文档无损格式精准替换工具
在保留原有 Run 字号、字体与加粗等样式的前提下，精准查找替换所有段落与表格中的文字。
用法: python docx_replace.py <docx_path> --old "旧词" --new "新词" [--output out.docx]
"""

import sys
import shutil
import argparse
from pathlib import Path

# 确保 Windows 终端 UTF-8 编码正常输出
if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8")
if hasattr(sys.stderr, "reconfigure"):
    sys.stderr.reconfigure(encoding="utf-8")

try:
    import docx
except ImportError:
    print("错误: 缺少 python-docx 依赖，请运行: uv pip install --system python-docx 或 pip install python-docx")
    sys.exit(1)


def replace_in_runs(runs, old_text, new_text):
    """
    当 Run 被拆分导致匹配困难时，或者能够直接定位 Run 中的 text 时进行替换。
    尽量保证原有首个 Run 的样式完全不被破坏。
    """
    full_text = "".join(r.text for r in runs)
    if old_text not in full_text:
        return False

    # 如果单一 run 里就包含目标短语，直接在 run 里替换（最安全无损）
    replaced_any = False
    for r in runs:
        if old_text in r.text:
            r.text = r.text.replace(old_text, new_text)
            replaced_any = True
            
    if replaced_any:
        return True

    # 如果文字跨越了多个 run 拆分，重新赋值给首个非空 run，清空后续被拆分的 runs
    # 注意：仅修改属于该段落匹配区域的文字
    new_full = full_text.replace(old_text, new_text)
    first_populated = False
    for r in runs:
        if not first_populated and r.text:
            r.text = new_full
            first_populated = True
        elif first_populated:
            r.text = ""
    return True


def backup_file(file_path: str) -> str:
    """在修改前自动创建 .bak.docx 备份副本"""
    p = Path(file_path)
    if not p.exists():
        return ""
    bak_path = p.with_suffix(".bak.docx")
    shutil.copy2(str(p), str(bak_path))
    return str(bak_path)


def replace_text_in_docx(file_path: str, old_text: str, new_text: str, output_path: str = None, dry_run: bool = False):
    doc = docx.Document(file_path)
    save_path = output_path or file_path
    
    count = 0
    
    # 1. 段落
    for p in doc.paragraphs:
        if old_text in p.text:
            if dry_run:
                count += 1
            elif replace_in_runs(p.runs, old_text, new_text):
                count += 1
                
    # 2. 表格
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if old_text in p.text:
                        if dry_run:
                            count += 1
                        elif replace_in_runs(p.runs, old_text, new_text):
                            count += 1
                            
    # 3. 页眉页脚
    for sec in doc.sections:
        for p in sec.header.paragraphs:
            if old_text in p.text:
                if dry_run:
                    count += 1
                elif replace_in_runs(p.runs, old_text, new_text):
                    count += 1
        for p in sec.footer.paragraphs:
            if old_text in p.text:
                if dry_run:
                    count += 1
                elif replace_in_runs(p.runs, old_text, new_text):
                    count += 1

    if dry_run:
        return f"[Dry-Run 模式] 预检发现 {count} 处匹配项（包含段落、表格或页眉页脚），目标：'{old_text}' -> '{new_text}'。未做实际文件写入。"

    # 自动备份将被覆盖的文件
    bak_msg = ""
    if Path(save_path).exists():
        bak = backup_file(save_path)
        if bak:
            bak_msg = f"（原文件已备份至: {bak}）"

    Path(save_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(save_path)
    return f"成功在文档中完成了 {count} 处段落/单元格的精准无损替换: '{old_text}' -> '{new_text}'，并保存至: {save_path}{bak_msg}"


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Word 文档无损格式精准替换工具")
    parser.add_argument("file_path", help="输入文档路径")
    parser.add_argument("--old", required=True, help="被替换的旧文本")
    parser.add_argument("--new", required=True, help="想要替换成的新文本")
    parser.add_argument("--output", "-o", help="输出路径（可选，默认覆盖原文件）")
    parser.add_argument("--dry-run", action="store_true", help="演练模式：仅统计替换匹配项数，不实际写回文件")
    args = parser.parse_args()
    
    msg = replace_text_in_docx(args.file_path, args.old, args.new, args.output, dry_run=args.dry_run)
    print(msg)

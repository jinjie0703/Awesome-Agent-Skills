#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
docx_template.py - Word 模板渲染填充工具
优先使用 docxtpl 进行 Jinja2 语法渲染；若未安装，则退化为 python-docx 基本变量占位符 {{ key }} 填充。
用法: python docx_template.py <template.docx> --json '{"name": "张三"}' --output out.docx
"""

import sys
import json
import shutil
import argparse
from pathlib import Path

# 确保 Windows 终端 UTF-8 编码正常输出
if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8")
if hasattr(sys.stderr, "reconfigure"):
    sys.stderr.reconfigure(encoding="utf-8")


def backup_file(file_path: str) -> str:
    """在修改前自动创建 .bak.docx 备份副本"""
    p = Path(file_path)
    if not p.exists():
        return ""
    bak_path = p.with_suffix(".bak.docx")
    shutil.copy2(str(p), str(bak_path))
    return str(bak_path)


def render_template(template_path: str, context: dict, output_path: str):
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    
    # 如果目标输出文件已存在，备份它
    bak_msg = ""
    if Path(output_path).exists():
        bak = backup_file(output_path)
        if bak:
            bak_msg = f"（原目标文件已备份至: {bak}）"

    # 优先尝试 docxtpl
    try:
        from docxtpl import DocxTemplate  # type: ignore
        doc = DocxTemplate(template_path)
        doc.render(context)
        doc.save(output_path)
        return f"使用 docxtpl 渲染成功，保存至: {output_path}{bak_msg}"
    except ImportError:
        pass

    # 退回使用 python-docx 基础填充 {{ key }}
    try:
        import docx
    except ImportError:
        return "错误: 既未安装 docxtpl，也未安装 python-docx。请运行: uv pip install --system docxtpl python-docx 或 pip install docxtpl python-docx"

    doc = docx.Document(template_path)
    
    def replace_placeholder_in_runs(runs, k, v):
        placeholder = f"{{{{ {k} }}}}"
        placeholder_tight = f"{{{{{k}}}}}"
        full = "".join(r.text for r in runs)
        if placeholder not in full and placeholder_tight not in full:
            return False
            
        new_full = full.replace(placeholder, str(v)).replace(placeholder_tight, str(v))
        first = True
        for r in runs:
            if first and r.text:
                r.text = new_full
                first = False
            elif not first:
                r.text = ""
        return True

    for p in doc.paragraphs:
        for k, v in context.items():
            replace_placeholder_in_runs(p.runs, k, v)
            
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for k, v in context.items():
                        replace_placeholder_in_runs(p.runs, k, v)
                        
    doc.save(output_path)
    return f"使用基础渲染完成 {{占位符}} 替换，保存至: {output_path}"


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Word 模板变量填充工具")
    parser.add_argument("template_path", help="Word 模板文件路径 (.docx)")
    group = parser.add_mutually_exclusive_group(required=True)
    group.add_argument("--json", dest="json_str", help="传入的 JSON 字符串参数")
    group.add_argument("--json-file", dest="json_file", help="传入的 JSON 配置文件路径（强烈推荐复杂内容或 Windows 命令行下使用）")
    parser.add_argument("--output", "-o", required=True, help="最终渲染生成的 Word 输出路径 (.docx)")
    args = parser.parse_args()
    
    try:
        if args.json_file:
            with open(args.json_file, "r", encoding="utf-8") as f:
                ctx = json.load(f)
        else:
            ctx = json.loads(args.json_str)
    except Exception as e:
        print(f"错误: JSON 解析失败 - {e}")
        sys.exit(1)
        
    msg = render_template(args.template_path, ctx, args.output)
    print(msg)

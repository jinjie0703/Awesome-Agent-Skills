#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
docx_format.py - Word 文档智能原子化排版引擎
支持标准预设 (thesis/official_doc) 以及用户完全定制化排版参数 (CLI/JSON)
用法:
  python docx_format.py input.docx --preset thesis --output formatted.docx
  python docx_format.py input.docx --body-font "宋体" --body-size 12 --line-spacing 1.5 --output out.docx
"""

import sys
import json
import shutil
import argparse
from pathlib import Path
try:
    import docx
    from docx.shared import Pt, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except ImportError:
    print("错误: 缺少 python-docx 依赖，请运行: uv pip install --system python-docx 或 pip install python-docx")
    sys.exit(1)


def apply_font_to_run(run, font_name_cn=None, font_name_ascii=None, size_pt=None, bold=None):
    """安全为 Run 赋予中英双字体及字号，自动处理 rPr 不存在的情况"""
    if font_name_ascii:
        run.font.name = font_name_ascii
    elif font_name_cn:
        run.font.name = font_name_cn
        
    if font_name_cn:
        # 安全获取或创建 rPr 节点，避免 rPr 为 None 时崩溃
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            from lxml import etree
            rFonts = etree.SubElement(rPr, qn('w:rFonts'))
        rFonts.set(qn('w:eastAsia'), font_name_cn)
        
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold


def _get_heading_level(style_name: str) -> int:
    """从样式名中提取标题层级，非标题返回 0"""
    if not style_name:
        return 0
    # 英文样式: "Heading 1", "Heading 2", ...
    if style_name.startswith("Heading "):
        try:
            return int(style_name.split(" ")[1])
        except (IndexError, ValueError):
            return 0
    # 中文样式: "标题 1", "标题 2", ...
    if style_name.startswith("标题 "):
        try:
            return int(style_name.split(" ")[1])
        except (IndexError, ValueError):
            return 0
    return 0


def _apply_heading_style(paragraph, config: dict, level: int):
    """根据标题层级应用对应的排版配置"""
    key_font = f"h{level}_font_cn"
    key_size = f"h{level}_size_pt"
    key_align = f"h{level}_align"
    
    # 对齐
    align_val = config.get(key_align)
    if align_val:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if align_val == "center" else WD_ALIGN_PARAGRAPH.LEFT
    
    # 字体与字号
    font_cn = config.get(key_font, config.get(f"h{level}_font"))
    size_pt = config.get(key_size, config.get(f"h{level}_size"))
    font_ascii = config.get("body_font_ascii", "Times New Roman")
    
    for r in paragraph.runs:
        apply_font_to_run(
            r,
            font_name_cn=font_cn,
            font_name_ascii=font_ascii,
            size_pt=size_pt,
            bold=True
        )


def _apply_body_style(paragraph, config: dict):
    """应用正文段落排版"""
    if "line_spacing" in config:
        paragraph.paragraph_format.line_spacing = config["line_spacing"]
    elif "line_spacing_pt" in config:
        paragraph.paragraph_format.line_spacing = Pt(config["line_spacing_pt"])
        
    if "first_line_indent_chars" in config:
        size_pt = config.get("body_size_pt", config.get("body_size", 12))
        paragraph.paragraph_format.first_line_indent = Pt(size_pt * config["first_line_indent_chars"])
        
    for r in paragraph.runs:
        apply_font_to_run(
            r,
            font_name_cn=config.get("body_font_cn", config.get("body_font")),
            font_name_ascii=config.get("body_font_ascii"),
            size_pt=config.get("body_size_pt", config.get("body_size"))
        )


def _format_paragraph(paragraph, config: dict):
    """统一处理单个段落（标题或正文）"""
    text = paragraph.text.strip()
    if not text:
        return
        
    style_name = paragraph.style.name if paragraph.style else "Normal"
    level = _get_heading_level(style_name)
    
    if level >= 1:
        _apply_heading_style(paragraph, config, level)
    else:
        _apply_body_style(paragraph, config)


def backup_file(file_path: str) -> str:
    """在修改前自动创建 .bak.docx 备份副本"""
    p = Path(file_path)
    bak_path = p.with_suffix(".bak.docx")
    shutil.copy2(str(p), str(bak_path))
    return str(bak_path)


def format_document(file_path: str, output_path: str, preset: str = None, custom_config: dict = None):
    doc = docx.Document(file_path)
    
    # 自动备份原文件
    bak = backup_file(file_path)
    
    # 整合排版配置
    config = {}
    
    # 1. 预设配置
    if preset == "thesis":
        config = {
            "page_margin_cm": 2.54,
            "body_font_cn": "宋体",
            "body_font_ascii": "Times New Roman",
            "body_size_pt": 12, # 小四
            "line_spacing": 1.5,
            "first_line_indent_chars": 2,
            "h1_font_cn": "黑体",
            "h1_size_pt": 18, # 二号
            "h1_align": "center",
            "h2_font_cn": "黑体",
            "h2_size_pt": 15, # 小三
            "h2_align": "left",
            "h3_font_cn": "黑体",
            "h3_size_pt": 14, # 四号
            "h3_align": "left",
            "h4_font_cn": "黑体",
            "h4_size_pt": 12, # 小四
            "h4_align": "left",
        }
    elif preset == "official_doc":
        config = {
            "page_margin_cm": 2.8,
            "body_font_cn": "仿宋_GB2312",
            "body_font_ascii": "Times New Roman",
            "body_size_pt": 16, # 三号
            "line_spacing_pt": 28, # 固定行距 28 磅
            "first_line_indent_chars": 2,
            "h1_font_cn": "方正小标宋简体",
            "h1_size_pt": 22, # 二号
            "h1_align": "center",
            "h2_font_cn": "黑体",
            "h2_size_pt": 16, # 三号
            "h2_align": "left",
            "h3_font_cn": "黑体",
            "h3_size_pt": 16, # 三号
            "h3_align": "left",
        }

    # 2. 用自定义配置覆盖（确保用户输入优先级最高）
    if custom_config:
        for k, v in custom_config.items():
            if v is not None:
                config[k] = v

    # 设置页面边距
    if "page_margin_cm" in config and doc.sections:
        for sec in doc.sections:
            sec.top_margin = Cm(config["page_margin_cm"])
            sec.bottom_margin = Cm(config["page_margin_cm"])
            sec.left_margin = Cm(config["page_margin_cm"] * 1.2) # 装订侧稍留宽
            sec.right_margin = Cm(config["page_margin_cm"])

    # 遍历并处理各段落（含 H1~H4 标题层级）
    for p in doc.paragraphs:
        _format_paragraph(p, config)

    # 遍历并处理表格内段落
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _format_paragraph(p, config)

    save_path = output_path or file_path
    Path(save_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(save_path)
    return f"文档排版完成并成功保存至: {save_path}（原文件已备份至: {bak}）"


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Word 文档排版引擎")
    parser.add_argument("file_path", help="输入文档路径")
    parser.add_argument("--output", "-o", help="输出路径（可选，默认覆盖原文件）")
    parser.add_argument("--preset", choices=["thesis", "official_doc"], help="应用标准预设")
    parser.add_argument("--custom-config", help="外部 JSON 排版配置文件路径")
    
    # CLI 细微参数控制
    parser.add_argument("--body-font", help="正文中文字体 (如 宋体)")
    parser.add_argument("--body-font-ascii", help="正文英文字体 (如 Times New Roman)")
    parser.add_argument("--body-size", type=float, help="正文字号 (磅值, 如 12)")
    parser.add_argument("--line-spacing", type=float, help="多倍行距 (如 1.5)")
    parser.add_argument("--first-line-indent", type=float, dest="first_line_indent_chars", help="首行缩进字符数 (如 2)")
    parser.add_argument("--h1-font", help="一级标题字体")
    parser.add_argument("--h1-size", type=float, help="一级标题字号")
    parser.add_argument("--h1-align", choices=["center", "left"], help="一级标题对齐方式")
    parser.add_argument("--h2-font", help="二级标题字体")
    parser.add_argument("--h2-size", type=float, help="二级标题字号")
    parser.add_argument("--h3-font", help="三级标题字体")
    parser.add_argument("--h3-size", type=float, help="三级标题字号")
    parser.add_argument("--h4-font", help="四级标题字体")
    parser.add_argument("--h4-size", type=float, help="四级标题字号")
    
    args = parser.parse_args()
    
    custom_dict = {}
    if args.custom_config and Path(args.custom_config).exists():
        with open(args.custom_config, "r", encoding="utf-8") as f:
            custom_dict = json.load(f)
            
    # 从命令行收集参数
    cli_mapping = {
        "body_font": "body_font_cn",
        "body_size": "body_size_pt",
        "h1_font": "h1_font_cn", "h1_size": "h1_size_pt",
        "h2_font": "h2_font_cn", "h2_size": "h2_size_pt",
        "h3_font": "h3_font_cn", "h3_size": "h3_size_pt",
        "h4_font": "h4_font_cn", "h4_size": "h4_size_pt",
    }
    for key in ["body_font", "body_font_ascii", "body_size", "line_spacing", "first_line_indent_chars",
                "h1_font", "h1_size", "h1_align",
                "h2_font", "h2_size",
                "h3_font", "h3_size",
                "h4_font", "h4_size"]:
        val = getattr(args, key, None)
        if val is not None:
            config_key = cli_mapping.get(key, key)
            custom_dict[config_key] = val

    result_msg = format_document(args.file_path, args.output, args.preset, custom_dict if custom_dict else None)
    print(result_msg)

#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
docx_inspect.py - Word 文档内容与排版侦查工具
用法: python docx_inspect.py <docx_path> [--preview-lines 100]
"""

import sys
import re
import json
import argparse
from pathlib import Path

# 确保 Windows 终端 UTF-8 编码正常输出
if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8")
if hasattr(sys.stderr, "reconfigure"):
    sys.stderr.reconfigure(encoding="utf-8")

try:
    import docx
    from docx.shared import Pt, Inches, Cm
except ImportError:
    print(json.dumps({"error": "缺少 python-docx 依赖，请运行: uv pip install --system python-docx 或 pip install python-docx"}))
    sys.exit(1)


def _get_heading_level(style_name: str) -> int:
    """从样式名中提取标题层级（支持 标题1, 标题 1, Heading1, Heading 1 等变体），非标题返回 0"""
    if not style_name:
        return 0
    match = re.match(r'^(?:heading|标题)\s*([1-6])$', style_name.strip(), re.IGNORECASE)
    return int(match.group(1)) if match else 0


def _collect_paragraph_info(paragraph, font_stats: dict, size_stats: dict):
    """从段落的 Runs 中统计字体与字号分布"""
    for run in paragraph.runs:
        if run.font:
            if run.font.name:
                font_stats[run.font.name] = font_stats.get(run.font.name, 0) + 1
            if run.font.size:
                pt_val = round(run.font.size.pt, 1)
                size_stats[str(pt_val)] = size_stats.get(str(pt_val), 0) + 1


def inspect_docx(file_path: str, preview_lines: int = 100):
    path = Path(file_path)
    if not path.exists():
        return {"error": f"文件不存在: {file_path}"}

    doc = docx.Document(path)
    
    # 统计信息
    p_count = len(doc.paragraphs)
    t_count = len(doc.tables)
    
    # 提取页面设置 (取第一个节)
    page_setup = {}
    if doc.sections:
        sec = doc.sections[0]
        try:
            page_setup = {
                "top_margin_cm": round(sec.top_margin.cm, 2) if sec.top_margin else None,
                "bottom_margin_cm": round(sec.bottom_margin.cm, 2) if sec.bottom_margin else None,
                "left_margin_cm": round(sec.left_margin.cm, 2) if sec.left_margin else None,
                "right_margin_cm": round(sec.right_margin.cm, 2) if sec.right_margin else None,
                "page_width_cm": round(sec.page_width.cm, 2) if sec.page_width else None,
                "page_height_cm": round(sec.page_height.cm, 2) if sec.page_height else None,
            }
        except Exception:
            page_setup = {"note": "提取页面边距失败"}

    # 提取样式与字体信息
    font_stats = {}
    size_stats = {}
    preview_md = []
    lines_collected = 0
    
    # 1. 普通段落
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
            
        # 判断标题与正文（支持 H1~H6）
        style_name = p.style.name if p.style else "Normal"
        level = _get_heading_level(style_name)
        prefix = "#" * level + " " if level >= 1 else ""
            
        if lines_collected < preview_lines:
            preview_md.append(f"{prefix}{text}")
            lines_collected += 1
            
        # 统计 Runs 的字体和字号
        _collect_paragraph_info(p, font_stats, size_stats)

    # 2. 表格内容（提取为 Markdown 表格简报）
    table_previews = []
    for t_idx, table in enumerate(doc.tables):
        rows_data = []
        for row in table.rows:
            cells_text = [cell.text.strip() for cell in row.cells]
            rows_data.append(cells_text)
            # 同时统计表格内字体分布
            for cell in row.cells:
                for p in cell.paragraphs:
                    _collect_paragraph_info(p, font_stats, size_stats)
        
        # 生成简单的 Markdown 表格预览
        if rows_data:
            header = rows_data[0]
            md_table_lines = [f"| {' | '.join(header)} |"]
            md_table_lines.append(f"| {' | '.join(['---'] * len(header))} |")
            for row_cells in rows_data[1:]:
                # 补齐列数不一致的情况
                padded = row_cells + [''] * (len(header) - len(row_cells))
                md_table_lines.append(f"| {' | '.join(padded[:len(header)])} |")
            table_previews.append({
                "table_index": t_idx + 1,
                "rows": len(rows_data),
                "cols": len(header),
                "preview_md": "\n".join(md_table_lines)
            })

    # 找出主要字体和主要字号
    main_font = max(font_stats, key=font_stats.get) if font_stats else "未知/默认"
    main_size_pt = max(size_stats, key=size_stats.get) if size_stats else "未知/默认"
    
    # 智能猜测文档类型倾向
    type_hint = "普通/自由文档 (建议保守排版)"
    full_text = "\n".join(preview_md)
    if "摘要" in full_text[:500] or "关键词" in full_text[:500] or "绪论" in full_text or "参考文献" in full_text:
        type_hint = "学术/毕业论文 (可选用 thesis 预设)"
    elif "关于" in full_text[:200] and ("通知" in full_text[:200] or "函" in full_text[:200] or "报告" in full_text[:200]):
        type_hint = "机关公文/函件 (可选用 official_doc 预设)"
    elif "证明" in full_text[:200] or "合同" in full_text[:200] or "简历" in full_text[:200]:
        type_hint = "表单/合同/证明书 (建议通过模板渲染或精准无损修改，不动排版)"

    return {
        "file_path": str(path.resolve()),
        "paragraph_count": p_count,
        "table_count": t_count,
        "document_type_hint": type_hint,
        "page_setup": page_setup,
        "style_summary": {
            "most_common_font": main_font,
            "most_common_size_pt": main_size_pt,
            "font_distribution": font_stats,
            "size_pt_distribution": size_stats
        },
        "content_preview_md": "\n".join(preview_md[:preview_lines]),
        "tables": table_previews
    }


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Word 文档内容与排版侦查工具")
    parser.add_argument("file_path", help="待检查的 .docx 文件路径")
    parser.add_argument("--preview-lines", type=int, default=100, help="预览行数限制")
    args = parser.parse_args()
    
    result = inspect_docx(args.file_path, args.preview_lines)
    print(json.dumps(result, ensure_ascii=False, indent=2))
